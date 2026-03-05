# -*- coding: utf-8 -*-
"""
Word 문서 자동 생성 모듈
회의록 포맷에 맞춰 문서 생성 (KH VATEC 양식)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime


def _set_cell_border(cell, **kwargs):
    """셀 테두리 설정"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text, font_name="맑은 고딕", font_size=10, bold=False, alignment=None):
    """셀 텍스트 설정"""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 셀 여백
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return run


def _set_row_height(row, height_cm):
    """행 높이 설정"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_paragraph(doc, text, font_name="맑은 고딕", font_size=10, bold=False,
                   left_indent=None, first_indent=None, space_before=0, space_after=0,
                   alignment=None, underline=False, color=None):
    """문단 추가 헬퍼"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def generate_meeting_minutes(meeting_info, transcript_text, output_path):
    """
    KH VATEC 양식 회의록 자동 생성
    
    Args:
        meeting_info (dict): 회의 정보
        transcript_text (str): 회의 내용 텍스트 (AI 또는 원본)
        output_path (str): 출력 파일 경로
    """
    doc = Document()
    
    # ===== 문서 여백 설정 =====
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # ===== 기본 스타일 설정 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # ===== 1. 제목: 회의록 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("회의록")
    title_run.font.name = '맑은 고딕'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.underline = True
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # ===== 2. 회의 정보 테이블 (4열) =====
    # 행1: 회의명 | 값 | 장소 | 값
    # 행2: 일시 | 값 | 작성자 | 값
    # 행3: 참석자 | 값 (3열 병합)
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 테이블 테두리 설정
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    
    # 전체 폭 설정
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '5000')
    tbl_width.set(qn('w:type'), 'pct')
    tblPr.append(tbl_width)
    
    # 열 너비: 라벨(좁은) / 값(넓은) / 라벨(좁은) / 값(넓은)
    col_widths = [Cm(2.5), Cm(6.0), Cm(2.5), Cm(5.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
    
    # --- 행 1: 회의명 / 장소 ---
    _set_row_height(table.rows[0], 0.8)
    _set_cell_shading(table.rows[0].cells[0], "D9D9D9")
    _set_cell_text(table.rows[0].cells[0], "회 의 명", bold=True, font_size=10)
    _set_cell_text(table.rows[0].cells[1], meeting_info.get('회의명', ''), font_size=10)
    _set_cell_shading(table.rows[0].cells[2], "D9D9D9")
    _set_cell_text(table.rows[0].cells[2], "장    소", bold=True, font_size=10)
    _set_cell_text(table.rows[0].cells[3], meeting_info.get('장소', ''), font_size=10)
    
    # --- 행 2: 일시 / 작성자 ---
    _set_row_height(table.rows[1], 0.8)
    _set_cell_shading(table.rows[1].cells[0], "D9D9D9")
    _set_cell_text(table.rows[1].cells[0], "일    시", bold=True, font_size=10)
    _set_cell_text(table.rows[1].cells[1], meeting_info.get('일시', ''), font_size=10)
    _set_cell_shading(table.rows[1].cells[2], "D9D9D9")
    _set_cell_text(table.rows[1].cells[2], "작 성 자", bold=True, font_size=10)
    _set_cell_text(table.rows[1].cells[3], meeting_info.get('작성자', ''), font_size=10)
    
    # --- 행 3: 참석자 (라벨 + 3열 병합) ---
    _set_row_height(table.rows[2], 0.8)
    _set_cell_shading(table.rows[2].cells[0], "D9D9D9")
    _set_cell_text(table.rows[2].cells[0], "참 석 자", bold=True, font_size=10)
    
    # 3열 병합 (cells[1] + cells[2] + cells[3])
    merged_cell = table.rows[2].cells[1].merge(table.rows[2].cells[3])
    attendees = meeting_info.get('참석자', '')
    company = meeting_info.get('업체이름', '')
    if company and company != '미정':
        attendees_text = f"{company}: {attendees}" if attendees else company
    else:
        attendees_text = attendees
    _set_cell_text(merged_cell, attendees_text, font_size=10)
    
    # ===== 3. "미 팅 내 용" 섹션 헤더 =====
    doc.add_paragraph()  # 빈 줄
    
    # 미팅 내용 헤더를 테이블로 (테두리 박스 효과)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl = header_table._tbl
    header_tblPr = header_tbl.tblPr if header_tbl.tblPr is not None else OxmlElement('w:tblPr')
    h_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    header_tblPr.append(h_borders)
    
    h_width = OxmlElement('w:tblW')
    h_width.set(qn('w:w'), '3000')
    h_width.set(qn('w:type'), 'pct')
    header_tblPr.append(h_width)
    
    header_cell = header_table.rows[0].cells[0]
    _set_cell_text(header_cell, "미 팅 내 용", font_size=12, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph()  # 빈 줄
    
    # ===== 4. 미팅 내용 본문 =====
    lines = transcript_text.strip().split('\n')
    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            # 빈 줄
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue
        
        # 들여쓰기 레벨 감지
        content = stripped.lstrip()
        leading_spaces = len(stripped) - len(content)
        
        # 굵은 텍스트 감지 (마크다운 ** 패턴)
        is_bold = False
        if content.startswith('**') and '**' in content[2:]:
            is_bold = True
            content = content.replace('**', '')
        
        # 들여쓰기 계산 (탭 또는 스페이스 기반)
        indent_cm = 0
        if leading_spaces >= 16 or content.startswith('→') or content.startswith('->'):
            indent_cm = 3.0
        elif leading_spaces >= 12 or (len(content) > 2 and content[0].isalpha() and content[1] == ')'):
            indent_cm = 2.4
        elif leading_spaces >= 8 and (content[:1].lower() in 'iv' or (len(content) > 1 and content[:2] in ['i.', 'ii', 'iv', 'v.'])):
            indent_cm = 1.8
        elif leading_spaces >= 4 or (len(content) > 2 and content[0].isdigit() and content[1] == ')'):
            indent_cm = 1.2
        elif content and content[0].isdigit() and '.' in content[:3]:
            indent_cm = 0.5
        
        # 번호 패턴별 자동 들여쓰기
        import re
        if re.match(r'^\d+\.\s', content):  # "1. " 패턴
            indent_cm = 0.5
            is_bold = True
        elif re.match(r'^\d+\)\s', content):  # "1) " 패턴
            indent_cm = 1.0
        elif re.match(r'^[ivxIVX]+\.\s', content):  # "i. " 패턴
            indent_cm = 1.5
        elif re.match(r'^[a-z]\)\s', content):  # "a) " 패턴
            indent_cm = 2.0
        elif content.startswith('→') or content.startswith('->'):  # 화살표
            indent_cm = 2.5
            content = content.replace('->', '→')
        
        # 문단 추가
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.left_indent = Cm(indent_cm)
        
        # 밑줄 텍스트 처리 (언더바 패턴)
        # 간단하게 텍스트로 추가
        run = p.add_run(content)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        run.font.bold = is_bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # ===== 5. 하단 생성 정보 =====
    doc.add_paragraph()  # 빈 줄
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(12)
    footer_run = footer_p.add_run(f"생성일시: {datetime.now().strftime('%Y.%m.%d %H:%M:%S')}")
    footer_run.font.name = '맑은 고딕'
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # ===== 저장 =====
    doc.save(output_path)
    print(f"✅ 문서 저장 완료: {output_path}")
    
    return output_path
